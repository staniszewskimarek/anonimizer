import io
import os
import httpx

from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from fastapi.responses import FileResponse, Response
from docx import Document

from anonymizer import anonymize_text, DEFAULT_MODEL

OLLAMA_TAGS_URL = "http://localhost:11434/api/tags"

app = FastAPI(title="Anonimizer PII")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))


@app.get("/")
def index():
    return FileResponse(os.path.join(BASE_DIR, "index.html"), media_type="text/html")


@app.get("/models")
async def get_models():
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            response = await client.get(OLLAMA_TAGS_URL)
            response.raise_for_status()
        return response.json()
    except Exception:
        raise HTTPException(status_code=503, detail="Nie można pobrać listy modeli z Ollamy.")


@app.post("/anonymize")
async def anonymize(
    file: UploadFile = File(...),
    model: str = Form(DEFAULT_MODEL),
):
    filename = file.filename or "output"
    ext = os.path.splitext(filename)[1].lower()

    if ext not in (".txt", ".docx"):
        raise HTTPException(status_code=400, detail="Obsługiwane formaty: TXT i DOCX.")

    content = await file.read()

    try:
        if ext == ".txt":
            text = content.decode("utf-8", errors="replace")
            result_text = anonymize_text(text, model)
            return Response(
                content=result_text.encode("utf-8"),
                media_type="text/plain; charset=utf-8",
                headers={"Content-Disposition": f'attachment; filename="{filename}"'},
            )

        else:  # .docx
            doc = Document(io.BytesIO(content))

            for para in doc.paragraphs:
                if para.text.strip():
                    anonymized = anonymize_text(para.text, model)
                    if para.runs:
                        para.runs[0].text = anonymized
                        for run in para.runs[1:]:
                            run.text = ""
                    else:
                        para.text = anonymized

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                anonymized = anonymize_text(para.text, model)
                                if para.runs:
                                    para.runs[0].text = anonymized
                                    for run in para.runs[1:]:
                                        run.text = ""
                                else:
                                    para.text = anonymized

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return Response(
                content=buf.read(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{filename}"'},
            )

    except httpx.ConnectError:
        raise HTTPException(
            status_code=503,
            detail="Nie można połączyć się z Ollama. Upewnij się, że serwer działa na localhost:11434.",
        )
    except httpx.TimeoutException:
        raise HTTPException(
            status_code=503,
            detail="Przekroczono limit czasu odpowiedzi Ollamy. Spróbuj z krótszym dokumentem.",
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Błąd serwera: {exc}")
